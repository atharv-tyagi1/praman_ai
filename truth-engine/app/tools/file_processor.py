"""
File processor — Extract text from uploaded documents, images, and audio files.
Supports PDF, DOCX, images (via Groq vision), and audio (via Groq Whisper).
"""

import io
import json
import base64
import logging
import tempfile
import os
import groq
from typing import Optional

from app.config import GROQ_API_KEY

# Configure Groq client
client = groq.Groq(api_key=GROQ_API_KEY)


async def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())
        
        full_text = "\n\n".join(text_parts)
        logging.info(f"PDF extraction: {len(reader.pages)} pages, {len(full_text)} chars")
        return full_text
    except Exception as e:
        logging.error(f"PDF extraction error: {e}")
        return f"[Error extracting PDF text: {str(e)}]"


async def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        full_text = "\n\n".join(text_parts)
        logging.info(f"DOCX extraction: {len(doc.paragraphs)} paragraphs, {len(full_text)} chars")
        return full_text
    except Exception as e:
        logging.error(f"DOCX extraction error: {e}")
        return f"[Error extracting DOCX text: {str(e)}]"


async def extract_text_from_image(file_bytes: bytes, mime_type: str = "image/jpeg") -> str:
    """
    Extract text/description from an image using Groq's vision model.
    Uses OCR-style prompting to get all visible text and context.
    """
    try:
        b64_data = base64.b64encode(file_bytes).decode("utf-8")
        
        completion = client.chat.completions.create(
            model="llama-3.2-11b-vision-preview",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "You are an expert text extractor and image analyst. "
                                "Analyze this image thoroughly and extract ALL visible text, "
                                "including headlines, body text, captions, labels, watermarks, "
                                "and any other textual content. Also describe the key visual "
                                "elements and context of the image.\n\n"
                                "Format your response as:\n"
                                "EXTRACTED TEXT:\n[all visible text]\n\n"
                                "IMAGE DESCRIPTION:\n[description of what the image shows]"
                            )
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{b64_data}"
                            }
                        }
                    ]
                }
            ],
            temperature=0.1,
            max_tokens=4096,
        )
        
        result = completion.choices[0].message.content.strip()
        logging.info(f"Image extraction: {len(result)} chars")
        return result
    except Exception as e:
        logging.error(f"Image extraction error: {e}")
        return f"[Error extracting text from image: {str(e)}]"


async def transcribe_audio(file_bytes: bytes, filename: str = "audio.mp3") -> str:
    """
    Transcribe audio using Groq's Whisper model.
    Supports: mp3, mp4, mpeg, mpga, m4a, wav, webm (max 25MB).
    """
    try:
        # Write to a temp file since Groq SDK needs a file-like object with name
        suffix = os.path.splitext(filename)[1] or ".mp3"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            with open(tmp_path, "rb") as audio_file:
                transcription = client.audio.transcriptions.create(
                    model="whisper-large-v3",
                    file=audio_file,
                    response_format="text",
                )
            
            result = str(transcription).strip()
            logging.info(f"Audio transcription: {len(result)} chars from {filename}")
            return result
        finally:
            os.unlink(tmp_path)
    except Exception as e:
        logging.error(f"Audio transcription error: {e}")
        return f"[Error transcribing audio: {str(e)}]"


# File type → handler mapping
SUPPORTED_EXTENSIONS = {
    # Documents
    ".pdf": "document",
    ".docx": "document",
    ".doc": "document",
    ".txt": "document",
    # Images
    ".jpg": "image",
    ".jpeg": "image",
    ".png": "image",
    ".gif": "image",
    ".webp": "image",
    ".bmp": "image",
    # Audio/Video
    ".mp3": "audio",
    ".wav": "audio",
    ".m4a": "audio",
    ".ogg": "audio",
    ".webm": "audio",
    ".mp4": "audio",
    ".mpeg": "audio",
    ".mpga": "audio",
}


async def process_uploaded_file(
    file_bytes: bytes,
    filename: str,
    content_type: Optional[str] = None,
    source_type: Optional[str] = None,
) -> dict:
    """
    Process an uploaded file and extract text content.
    
    Args:
        file_bytes: Raw file bytes
        filename: Original filename
        content_type: MIME type of the file
        source_type: Hint from frontend ('document', 'image', 'audio')
    
    Returns:
        dict with 'text', 'source_info', and optional 'error'
    """
    ext = os.path.splitext(filename)[1].lower()
    
    # Determine file type
    file_type = source_type or SUPPORTED_EXTENSIONS.get(ext, "unknown")
    
    result = {
        "text": "",
        "source_info": {
            "filename": filename,
            "file_type": file_type,
            "file_size": len(file_bytes),
            "content_type": content_type,
        },
        "error": None,
    }
    
    try:
        if file_type == "document":
            if ext == ".pdf":
                result["text"] = await extract_text_from_pdf(file_bytes)
            elif ext in (".docx", ".doc"):
                result["text"] = await extract_text_from_docx(file_bytes)
            elif ext == ".txt":
                result["text"] = file_bytes.decode("utf-8", errors="replace")
            else:
                result["error"] = f"Unsupported document format: {ext}"
        
        elif file_type == "image":
            mime = content_type or f"image/{ext.lstrip('.')}"
            result["text"] = await extract_text_from_image(file_bytes, mime)
        
        elif file_type == "audio":
            result["text"] = await transcribe_audio(file_bytes, filename)
        
        else:
            result["error"] = f"Unsupported file type: {ext}"
    
    except Exception as e:
        result["error"] = str(e)
        logging.error(f"File processing error for {filename}: {e}")
    
    return result
